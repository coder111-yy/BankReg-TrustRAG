from __future__ import annotations

import hashlib
import re
import subprocess
import shutil
import tempfile
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Any

from ..schemas import Document, TableCellEvidence, TextEvidence
from ..utils import (
    compact_path,
    first_nonempty,
    insurance_company_scope,
    is_insurance_fund_table,
    normalize_text,
    normalized_reporting_period,
    sha256_file,
    stable_id,
)


@dataclass
class ParseResult:
    document: Document
    text_evidence: list[TextEvidence] = field(default_factory=list)
    table_evidence: list[TableCellEvidence] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

def _is_percent_format(number_format: Any) -> bool:
    text = str(number_format or "")
    return "%" in text or "％" in text


def _apply_excel_percent_formats(
    records: list[TableCellEvidence],
    percent_cells: set[str],
) -> list[TableCellEvidence]:
    """
    Convert Excel percentage-formatted numeric cells from their stored
    decimal representation to percentage-point representation.

    Example:
        Excel stored value: 1.54729
        Excel display:      154.729%
        evidence value:     154.729
        evidence unit:      %
    """
    normalized: list[TableCellEvidence] = []

    for record in records:
        value = record.value

        if (
            record.cell_address in percent_cells
            and isinstance(value, (int, float))
            and not isinstance(value, bool)
        ):
            scaled = round(float(value) * 100.0, 12)

            # context 里面最后一个字段也是原始数值，
            # 必须同步修改，否则后面的 BGE / LLM / Verification
            # 仍可能看到 1.54729 而不是 154.729。
            context = normalize_text(record.context)

            if context:
                parts = [part.strip() for part in context.split("|")]
                if parts:
                    parts[-1] = str(scaled)
                    context = " | ".join(parts)

            # TableCellEvidence 不是 Pydantic model，
            # 使用 dataclasses.replace 创建修改后的副本。
            record = replace(
                record,
                value=scaled,
                unit="%",
                context=context,
            )

        normalized.append(record)

    return normalized
def _document_type(path: Path) -> str:
    suffix = path.suffix.lower()
    return {".docx": "word", ".doc": "word", ".pdf": "pdf", ".xlsx": "excel", ".xls": "excel"}.get(suffix, "other")


def _title(path: Path) -> str:
    stem = path.stem
    # Attachment names often repeat the title after an ordinal prefix and underscore.
    stem = re.sub(r"^\d+_", "", stem)
    parts = stem.split("_")
    if len(parts) > 1 and parts[-1] == parts[-2]:
        parts = parts[:-1]
    return normalize_text("_".join(parts))


def _date_from_name(text: str) -> str | None:
    """Extract a publication-like date conservatively.

    A year range such as ``2013年至2017年6月`` describes statistical
    coverage, not the publication date.  The old implementation converted the
    first year into ``2013-01-01``, which created false document metadata.
    """
    normalized = normalize_text(text)
    full = re.search(r"(20\d{2})年(\d{1,2})月(\d{1,2})日", normalized)
    if full:
        return f"{int(full.group(1)):04d}-{int(full.group(2)):02d}-{int(full.group(3)):02d}"

    # Do not interpret coverage ranges as publication dates.
    if re.search(r"20\d{2}年.{0,4}(?:至|到|—|-)\s*20\d{2}年", normalized):
        return None

    month = re.search(r"(20\d{2})年(\d{1,2})月", normalized)
    if month:
        return f"{int(month.group(1)):04d}-{int(month.group(2)):02d}-01"
    iso = re.search(r"(20\d{2})[-/](\d{1,2})[-/](\d{1,2})", normalized)
    if iso:
        return f"{int(iso.group(1)):04d}-{int(iso.group(2)):02d}-{int(iso.group(3)):02d}"
    return None


def _explicit_authority(text: str) -> str | None:
    """Return an authority only when the organization is explicitly named."""
    value = normalize_text(text)
    aliases = [
        ("国家金融监督管理总局", "国家金融监督管理总局"),
        ("中国银行保险监督管理委员会", "中国银行保险监督管理委员会"),
        ("中国银保监会", "中国银行保险监督管理委员会"),
        ("中国银行业监督管理委员会", "中国银行业监督管理委员会"),
        ("中国银监会", "中国银行业监督管理委员会"),
        ("中国保险监督管理委员会", "中国保险监督管理委员会"),
        ("中国保监会", "中国保险监督管理委员会"),
        ("中国人民银行", "中国人民银行"),
    ]
    for alias, canonical in aliases:
        if alias in value:
            return canonical
    return None


def _metadata(path: Path, root: Path) -> Document:
    title = _title(path)
    digest = sha256_file(path)

    # One physical source file must have one stable, unique doc_id even when
    # two files contain identical bytes. Content SHA remains available for
    # duplicate detection, while doc_id is derived from both relative path and
    # content hash.
    relative_path = compact_path(path, root).replace("\\", "/")
    identity = f"{relative_path}\0{digest}".encode("utf-8")
    doc_id = "DOC_" + hashlib.sha256(identity).hexdigest()[:16]

    date = _date_from_name(path.name)
    status = "unknown"
    return Document(
        doc_id=doc_id,
        title=title,
        authority=_explicit_authority(title),
        document_no=None,
        publish_date=date,
        effective_date=None,
        expire_date=None,
        document_type=_document_type(path),
        topic=[],
        version=None,
        status=status,
        source_url=None,
        local_path=compact_path(path, root),
        sha256=digest,
        file_name=path.name,
    )


_CN_NUM = r"一二三四五六七八九十百千万零〇0-9"


def _heading_kind(text: str) -> tuple[str | None, str | None]:
    """Classify common Chinese regulatory heading forms."""
    value = normalize_text(text)
    if not value:
        return None, None
    match = re.match(rf"^(第[{_CN_NUM}]+章)\s*(.*)$", value)
    if match:
        return "chapter", value
    match = re.match(rf"^(第[{_CN_NUM}]+节)\s*(.*)$", value)
    if match:
        return "section", value
    match = re.match(rf"^(第[{_CN_NUM}]+条)\s*(.*)$", value)
    if match:
        return "article", match.group(1)
    # Many explanatory/regulatory PDFs use 一、/（一） as section headings.
    if re.match(r"^[一二三四五六七八九十]+、\S+", value) and len(value) <= 80:
        return "section", value
    return None, None


def _paragraph_evidence(
    doc: Document,
    paragraphs: list[tuple[str, int | None, str | None, str | None, str | None]],
) -> list[TextEvidence]:
    evidence: list[TextEvidence] = []
    for index, (content, page, chapter, section, article_no) in enumerate(paragraphs, 1):
        content = normalize_text(content)
        if not content:
            continue
        evidence.append(
            TextEvidence(
                evidence_id=f"text:{doc.doc_id}:p{index}",
                doc_id=doc.doc_id,
                content=content,
                page=page,
                chapter=chapter,
                article_no=article_no,
                paragraph_no=index,
                section=section,
                source_url=doc.source_url,
                source_location=f"{doc.file_name}:page:{page or 'na'}:paragraph:{index}",
            )
        )
    return evidence


def _structured_paragraphs(
    blocks: list[tuple[str, int | None]],
) -> list[tuple[str, int | None, str | None, str | None, str | None]]:
    """Attach chapter/section/article scope to an ordered text stream."""
    result: list[tuple[str, int | None, str | None, str | None, str | None]] = []
    chapter: str | None = None
    section: str | None = None
    article: str | None = None
    for raw, page in blocks:
        content = normalize_text(raw)
        if not content:
            continue
        kind, label = _heading_kind(content)
        if kind == "chapter":
            chapter = content
            section = None
            article = None
        elif kind == "section":
            section = content
            article = None
        elif kind == "article":
            article = label
        result.append((content, page, chapter, section, article))
    return result


def _join_pdf_lines(text: str) -> list[str]:
    """Reconstruct readable PDF paragraphs without OCR.

    pypdf frequently returns one visual line at a time.  Joining those lines
    prevents a sentence such as ``关于报送绿色信 / 贷统计表的通知`` from
    becoming two unrelated evidence records.
    """
    lines = [normalize_text(line) for line in str(text or "").splitlines()]
    lines = [line for line in lines if line]
    output: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            output.append(normalize_text("".join(buffer)))
            buffer.clear()

    for line in lines:
        kind, _ = _heading_kind(line)
        if kind in {"chapter", "section", "article"}:
            flush()
            output.append(line)
            continue
        buffer.append(line)
        # Regulatory prose normally ends a semantic paragraph at Chinese/full
        # punctuation.  A colon is intentionally not a hard stop because lists
        # often continue on the next visual line.
        if re.search(r"[。！？；;]$", line):
            flush()
    flush()
    return output


def parse_docx(path: Path, doc: Document) -> list[TextEvidence]:
    from docx import Document as WordDocument

    blocks: list[tuple[str, int | None]] = []
    document = WordDocument(str(path))
    for paragraph in document.paragraphs:
        content = normalize_text(paragraph.text)
        if content:
            blocks.append((content, None))
    # Keep Word tables as traceable textual evidence.  Excel attachments still
    # use the dedicated structured table parser below.
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = " | ".join(normalize_text(cell.text) for cell in row.cells)
            if normalize_text(text):
                blocks.append((f"表{table_index} 行{row_index}: {text}", None))
    return _paragraph_evidence(doc, _structured_paragraphs(blocks))


def parse_legacy_doc(
    path: Path,
    doc: Document
) -> tuple[list[TextEvidence], list[str]]:
    """
    使用 LibreOffice 将旧版 .doc 转换为 .docx，
    然后复用现有的 parse_docx() 解析正文和表格。
    """

    # 1. 优先从 PATH 查找 soffice
    soffice = shutil.which("soffice")

    # 2. Windows 常见 LibreOffice 安装位置
    if not soffice:
        candidates = [
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        ]

        for candidate in candidates:
            if candidate.exists():
                soffice = str(candidate)
                break

    if not soffice:
        recovered = _parse_wps_utf16_doc(path)
        if recovered:
            blocks = [(line, None) for line in recovered.splitlines() if normalize_text(line)]
            return _paragraph_evidence(doc, _structured_paragraphs(blocks)), [
                "LibreOffice not found; recovered text from WPS WordDocument stream."
            ]
        return [], [
            "LibreOffice not found and WPS fallback recovered no usable text."
        ]

    try:
        # 使用临时目录，避免转换后的 docx 污染原始数据目录
        with tempfile.TemporaryDirectory(prefix="bankreg_doc_") as temp_dir:
            temp_path = Path(temp_dir)

            completed = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(temp_path),
                    str(path),
                ],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
            )

            if completed.returncode != 0:
                recovered = _parse_wps_utf16_doc(path)
                if recovered:
                    blocks = [(line, None) for line in recovered.splitlines() if normalize_text(line)]
                    return _paragraph_evidence(doc, _structured_paragraphs(blocks)), [
                        f"LibreOffice conversion failed (returncode={completed.returncode}); used WPS text fallback."
                    ]
                return [], [
                    f"LibreOffice conversion failed "
                    f"(returncode={completed.returncode}): "
                    f"{completed.stderr[:500]}"
                ]

            # LibreOffice 转换后的文件
            converted_path = temp_path / f"{path.stem}.docx"

            if not converted_path.exists():
                return [], [
                    f"LibreOffice conversion produced no DOCX file: {path}"
                ]

            # 直接复用现有 DOCX 解析逻辑
            evidence = parse_docx(converted_path, doc)

            return evidence, []

    except subprocess.TimeoutExpired:
        return [], [
            f"LibreOffice conversion timeout: {path}"
        ]

    except Exception as exc:
        return [], [
            f"LibreOffice .doc parser failed: {exc}"
        ]


def _parse_wps_utf16_doc(path: Path) -> str:
    """Recover text from WPS-generated OLE .doc files.

    Several contest attachments have a valid WordDocument stream but report
    zero words and are rejected by antiword.  WPS stores the visible body as
    UTF-16LE runs in that stream.  This conservative fallback extracts only
    printable CJK/Latin runs and leaves the original file untouched.
    """
    try:
        import olefile

        with olefile.OleFileIO(str(path)) as ole:
            if not ole.exists("WordDocument"):
                return ""
            raw = ole.openstream("WordDocument").read()
        decoded = raw.decode("utf-16le", errors="ignore")
        runs: list[str] = []
        current: list[str] = []

        def flush() -> None:
            if not current:
                return
            value = normalize_text("".join(current))
            if value and value not in {"PAGE", "MERGEFORMAT"}:
                runs.append(value)
            current.clear()

        for char in decoded:
            code = ord(char)
            if char in "\r\n\t" or 0x20 <= code <= 0x7E or 0x3000 <= code <= 0x9FFF:
                current.append(char)
            else:
                flush()
        flush()
        # Metadata and field names are short; the body contains several CJK
        # characters.  This also prevents returning an apparently successful
        # evidence record for an empty or unsupported .doc file.
        body = [run for run in runs if sum(0x4E00 <= ord(char) <= 0x9FFF for char in run) >= 2]
        return "\n".join(body)
    except (ImportError, OSError, ValueError):
        return ""


def parse_pdf(path: Path, doc: Document) -> list[TextEvidence]:
    from pypdf import PdfReader

    blocks: list[tuple[str, int | None]] = []
    reader = PdfReader(str(path))
    for page_number, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        blocks.extend((block, page_number) for block in _join_pdf_lines(text))
    return _paragraph_evidence(doc, _structured_paragraphs(blocks))


def _as_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (str, int, float, bool)):
        return value
    return normalize_text(value)


def _period_from_header(header: str | None) -> str | None:
    if not header:
        return None
    value = normalize_text(header)
    parsed_period = normalized_reporting_period(value)
    if parsed_period:
        return parsed_period
    m = re.search(r"(20\d{2})[-/]0?(\d{1,2})", value)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}"
    return value if re.search(r"\d", value) else None


_QUARTER_LABELS = {"一季度", "二季度", "三季度", "四季度"}
_MONTH_RE = re.compile(r"^(1[0-2]|0?[1-9])月$")


def _year_from_text(*values: Any) -> str | None:
    for value in values:
        match = re.search(r"(?<!\d)(20\d{2})(?:年)?(?!\d)", normalize_text(value))
        if match:
            return match.group(1)
    return None


def _month_label(value: Any) -> int | None:
    match = _MONTH_RE.fullmatch(normalize_text(value))
    return int(match.group(1)) if match else None


def _is_month_header_row(row: list[Any]) -> bool:
    months = [_month_label(value) for value in row]
    return sum(month is not None for month in months) >= 2


def _looks_like_month_blocks(rows: list[list[Any]]) -> bool:
    return sum(1 for row in rows[:120] if _is_month_header_row(row)) >= 1


def _quarter_label(value: Any) -> str | None:
    """Normalize the quarter marker used by block-form statistical sheets."""
    text = normalize_text(value)
    if text in _QUARTER_LABELS:
        return text
    match = re.fullmatch(r"第?([一二三四1-4])季度", text)
    if not match:
        return None
    return {
        "一": "一季度", "1": "一季度",
        "二": "二季度", "2": "二季度",
        "三": "三季度", "3": "三季度",
        "四": "四季度", "4": "四季度",
    }[match.group(1)]


def _infer_unit(*values: Any) -> str | None:
    match = re.search(r"(百分比|%|％|‰|亿元|万元|百万元|元)", " ".join(normalize_text(value) for value in values if value))
    return match.group(1) if match else None


def _infer_sheet_unit(rows: list[list[Any]], header_rows: int) -> str | None:
    """Return a sheet-wide unit only when the header is unambiguous.

    Some statistical workbooks place ``单位：亿元`` in a visually merged
    cell above just one physical column.  Per-column header extraction then
    loses the unit for neighbouring numeric cells.  Do not propagate when a
    header also advertises count-like units because those sheets are mixed.
    """
    header = " ".join(
        normalize_text(value)
        for row in rows[:header_rows]
        for value in row
        if normalize_text(value)
    )
    if re.search(r"(?:万件|件|万户|户|家|个)", header):
        return None
    units = {
        "%" if value in {"百分比", "%", "％"} else value
        for value in re.findall(r"百分比|%|％|‰|万亿元|亿元|百万元|万元|元", header)
    }
    return next(iter(units)) if len(units) == 1 else None


def _looks_like_quarter_matrix(rows: list[list[Any]], header_rows: int) -> bool:
    """Detect quarter blocks x indicator rows x institution columns."""
    if header_rows < 2 or max((len(row) for row in rows), default=0) < 3:
        return False
    active_quarter: str | None = None
    quarter_count = 0
    metric_rows = 0
    for row in rows[header_rows:header_rows + 60]:
        if not row:
            continue
        quarter = _quarter_label(row[0] if len(row) > 0 else None)
        if quarter:
            active_quarter = quarter
            quarter_count += 1
        metric = normalize_text(row[1] if len(row) > 1 else None)
        values = [value for value in row[2:] if value is not None and normalize_text(value)]
        if active_quarter and metric and values:
            metric_rows += 1
    return quarter_count >= 2 and metric_rows >= 4


def _quarter_matrix_records(
    doc: Document,
    sheet_name: str,
    rows: list[list[Any]],
    header_rows: int,
    headers: list[str],
    table_name: str | None,
) -> list[TableCellEvidence]:
    records: list[TableCellEvidence] = []
    year = _year_from_text(doc.title)
    sheet_unit = _infer_unit(" ".join(headers), " ".join(normalize_text(value) for row in rows[:header_rows] for value in row if value))
    active_quarter: str | None = None
    for row_index, row in enumerate(rows[header_rows:], header_rows + 1):
        quarter = _quarter_label(row[0] if len(row) > 0 else None)
        if quarter:
            active_quarter = quarter
        indicator = normalize_text(row[1] if len(row) > 1 else None) or None
        if not active_quarter or not indicator:
            continue
        for column_index, value in enumerate(row[2:], 3):
            if value is None or normalize_text(value) == "":
                continue
            column_header = headers[column_index - 1] if column_index - 1 < len(headers) else None
            column_header = column_header or None
            normalized_value = normalize_text(value)
            context = f"{indicator} | {active_quarter} | {column_header or ''} | {normalized_value}"
            address = _excel_address(row_index, column_index)
            records.append(
                TableCellEvidence(
                    evidence_id=f"cell:{doc.doc_id}:{sheet_name}:{address}",
                    doc_id=doc.doc_id,
                    sheet_name=sheet_name,
                    table_name=table_name,
                    indicator=indicator,
                    period=f"{year}年{active_quarter}" if year else active_quarter,
                    value=_as_value(value),
                    unit=_infer_unit(column_header, indicator) or sheet_unit,
                    row_header=indicator,
                    column_header=" / ".join(part for part in [active_quarter, column_header] if part),
                    cell_address=address,
                    context=normalize_text(context),
                    source_url=doc.source_url,
                )
            )
    return records


def _monthly_block_records(
    doc: Document,
    sheet_name: str,
    rows: list[list[Any]],
    table_name: str | None,
) -> list[TableCellEvidence]:
    """Parse repeated ``scope -> unit/year -> month header -> data`` blocks."""
    records: list[TableCellEvidence] = []
    active_scope: str | None = None
    active_unit: str | None = None
    active_year: str | None = _year_from_text(doc.title)
    month_columns: dict[int, int] = {}

    for row_index, row in enumerate(rows, 1):
        nonempty = [normalize_text(value) for value in row if normalize_text(value)]
        first = normalize_text(row[0] if row else None)

        if not nonempty:
            continue

        # Notes are valuable evidence but should not be mistaken for numeric cells.
        if first.startswith(("注：", "注:", "注释", "说明：", "说明:")):
            records.append(
                TableCellEvidence(
                    evidence_id=f"cell:{doc.doc_id}:{sheet_name}:{_excel_address(row_index, 1)}",
                    doc_id=doc.doc_id,
                    sheet_name=sheet_name,
                    table_name=table_name,
                    indicator="注",
                    period=active_year,
                    value=_as_value(first),
                    unit=None,
                    row_header="注",
                    column_header=active_scope,
                    cell_address=_excel_address(row_index, 1),
                    context=normalize_text(" | ".join(part for part in [active_scope, first] if part)),
                    source_url=doc.source_url,
                )
            )
            continue

        unit_text = " ".join(nonempty)
        if re.search(r"单位[:：]", unit_text):
            active_unit = _infer_unit(unit_text)
            continue

        # A dedicated year row often sits directly above the month header.
        row_year = _year_from_text(*nonempty)
        if row_year and (first in {"时间", "年份", "年度"} or len(nonempty) <= 3):
            active_year = row_year

        if _is_month_header_row(row):
            month_columns = {
                index: month
                for index, value in enumerate(row, 1)
                if (month := _month_label(value)) is not None
            }
            continue

        numeric_values = [
            value for value in row[1:]
            if isinstance(value, (int, float)) and not isinstance(value, bool)
        ]

        # Single-cell labels between blocks are institution/scope headings.
        if len(nonempty) == 1 and not numeric_values and not first.startswith(("时间", "项目")):
            if not re.search(r"^(?:20\d{2}年|单位[:：])", first):
                active_scope = first
            continue

        if not month_columns or not first:
            continue

        indicator = first
        percent_metric = bool(re.search(r"率|比例|占比|增幅|增长", indicator))
        for column_index, month in month_columns.items():
            if column_index > len(row):
                continue
            value = row[column_index - 1]
            if value is None or normalize_text(value) == "":
                continue
            # Month-block business cells are expected to be numeric.  Text cells
            # remain available through raw-cell export in manifest.py.
            if not isinstance(value, (int, float)) or isinstance(value, bool):
                continue
            period = f"{active_year}-{month:02d}" if active_year else f"{month}月"
            unit = "%" if percent_metric else (active_unit or _infer_unit(indicator))
            column_header = " / ".join(part for part in [active_scope, f"{month}月"] if part)
            address = _excel_address(row_index, column_index)
            context = " | ".join(
                str(part) for part in [active_scope, indicator, period, unit, normalize_text(value)] if part
            )
            records.append(
                TableCellEvidence(
                    evidence_id=f"cell:{doc.doc_id}:{sheet_name}:{address}",
                    doc_id=doc.doc_id,
                    sheet_name=sheet_name,
                    table_name=table_name,
                    indicator=indicator,
                    period=period,
                    value=_as_value(value),
                    unit=unit,
                    row_header=indicator,
                    column_header=column_header or f"{month}月",
                    cell_address=address,
                    context=normalize_text(context),
                    source_url=doc.source_url,
                )
            )
    return records


def _cell_records(doc: Document, sheet_name: str, rows: list[list[Any]], table_name: str | None = None) -> list[TableCellEvidence]:
    records: list[TableCellEvidence] = []
    if not rows:
        return records
    if _looks_like_month_blocks(rows):
        return _monthly_block_records(doc, sheet_name, rows, table_name)
    width = max((len(r) for r in rows), default=0)
    header_rows = _header_row_count(rows)
    headers: list[str] = []
    for column in range(width):
        pieces = [normalize_text(rows[row][column]) for row in range(header_rows) if column < len(rows[row]) and normalize_text(rows[row][column])]
        headers.append(" / ".join(dict.fromkeys(pieces)))
    if _looks_like_quarter_matrix(rows, header_rows):
        return _quarter_matrix_records(doc, sheet_name, rows, header_rows, headers, table_name)
    scoped_insurance_report = is_insurance_fund_table(doc.title, sheet_name, table_name)
    sheet_unit = _infer_sheet_unit(rows, header_rows)
    active_scope = "保险业总体" if scoped_insurance_report else None
    for row_index, row in enumerate(rows, 1):
        row_header = first_nonempty(row[:2])
        indicator = row_header
        if scoped_insurance_report:
            active_scope = insurance_company_scope(row_header) or active_scope
        for column_index, value in enumerate(row, 1):
            if value is None or normalize_text(value) == "":
                continue
            column_header = headers[column_index - 1] or None
            period = _period_from_header(column_header) or _period_from_header(doc.title)
            unit = None
            unit = _infer_unit(column_header, row_header)
            if unit is None and isinstance(value, (int, float)) and not isinstance(value, bool):
                unit = sheet_unit
            address = _excel_address(row_index, column_index)
            context_parts = [active_scope, indicator, column_header, normalize_text(value)]
            context = " | ".join(str(part) for part in context_parts if part)
            records.append(
                TableCellEvidence(
                    evidence_id=f"cell:{doc.doc_id}:{sheet_name}:{address}",
                    doc_id=doc.doc_id,
                    sheet_name=sheet_name,
                    table_name=table_name,
                    indicator=indicator,
                    period=period,
                    value=_as_value(value),
                    unit=unit,
                    row_header=row_header,
                    column_header=column_header,
                    cell_address=address,
                    context=normalize_text(context),
                    source_url=doc.source_url,
                )
            )
    return records


def _header_row_count(rows: list[list[Any]]) -> int:
    """Find the first data row instead of assuming every workbook has 4 headers.

    The old fixed four-row window included the first data row in column
    headers for sheets whose data starts on row 4.  That polluted headers and
    periods (for example ``合计 / 45167.98``), making exact row/column lookup
    impossible even though the source cell was present.
    """
    for row_index, row in enumerate(rows[:12]):
        if not row or not first_nonempty(row[:2]):
            continue
        numeric_values = [
            value for value in row[1:]
            if isinstance(value, (int, float)) and not isinstance(value, bool)
        ]
        if numeric_values:
            return max(1, row_index)
    return min(4, len(rows))


def _excel_address(row: int, column: int) -> str:
    result = ""
    number = column
    while number:
        number, remainder = divmod(number - 1, 26)
        result = chr(65 + remainder) + result
    return f"{result}{row}"


def _parse_xls_with_libreoffice(path: Path, doc: Document) -> tuple[list[TableCellEvidence], list[str]]:
    """Fallback for legacy XLS when xlrd is unavailable."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return [], ["xlrd is unavailable and LibreOffice was not found for XLS fallback."]
    try:
        with tempfile.TemporaryDirectory(prefix="bankreg_xls_") as temp_dir:
            completed = subprocess.run(
                [soffice, "--headless", "--convert-to", "xlsx", "--outdir", temp_dir, str(path)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
            )
            converted = Path(temp_dir) / f"{path.stem}.xlsx"
            if completed.returncode != 0 or not converted.exists():
                return [], [f"LibreOffice XLS conversion failed: {completed.stderr[:500]}"]
            import openpyxl

            workbook = openpyxl.load_workbook(converted, read_only=True, data_only=True)
            records: list[TableCellEvidence] = []
            for sheet in workbook.worksheets:
                rows = [list(row) for row in sheet.iter_rows(values_only=True)]
                records.extend(_cell_records(doc, sheet.title, rows, sheet.title))
            return records, ["legacy XLS parsed through LibreOffice XLSX fallback"]
    except Exception as exc:
        return [], [f"legacy XLS fallback failed: {type(exc).__name__}: {exc}"]


def parse_excel(
    path: Path,
    doc: Document,
) -> tuple[list[TableCellEvidence], list[str]]:
    warnings: list[str] = []
    records: list[TableCellEvidence] = []

    try:
        # ============================================================
        # XLSX
        # ============================================================
        if path.suffix.lower() == ".xlsx":
            import openpyxl

            workbook = openpyxl.load_workbook(
                path,
                read_only=True,
                data_only=True,
            )

            for sheet in workbook.worksheets:
                rows: list[list[Any]] = []
                percent_cells: set[str] = set()

                # 不再 values_only=True，
                # 因为必须保留 number_format。
                for excel_row in sheet.iter_rows():
                    values: list[Any] = []

                    for cell in excel_row:
                        values.append(cell.value)

                        if _is_percent_format(
                            getattr(cell, "number_format", None)
                        ):
                            percent_cells.add(cell.coordinate)

                    rows.append(values)

                sheet_records = _cell_records(
                    doc,
                    sheet.title,
                    rows,
                    sheet.title,
                )

                records.extend(
                    _apply_excel_percent_formats(
                        sheet_records,
                        percent_cells,
                    )
                )

        # ============================================================
        # XLS
        # ============================================================
        else:
            import xlrd

            workbook = xlrd.open_workbook(
                str(path),
                on_demand=True,
                formatting_info=True,
            )

            for sheet in workbook.sheets():
                rows: list[list[Any]] = []
                percent_cells: set[str] = set()

                for row_index in range(sheet.nrows):
                    values: list[Any] = []

                    for column_index in range(sheet.ncols):
                        cell = sheet.cell(
                            row_index,
                            column_index,
                        )

                        values.append(cell.value)

                        try:
                            xf = workbook.xf_list[cell.xf_index]
                            fmt = workbook.format_map.get(
                                xf.format_key
                            )
                            format_string = (
                                fmt.format_str
                                if fmt is not None
                                else ""
                            )

                            if _is_percent_format(format_string):
                                percent_cells.add(
                                    _excel_address(
                                        row_index + 1,
                                        column_index + 1,
                                    )
                                )
                        except (
                            AttributeError,
                            IndexError,
                            KeyError,
                            TypeError,
                        ):
                            pass

                    rows.append(values)

                sheet_records = _cell_records(
                    doc,
                    sheet.name,
                    rows,
                    sheet.name,
                )

                records.extend(
                    _apply_excel_percent_formats(
                        sheet_records,
                        percent_cells,
                    )
                )

    except ImportError as exc:
        warnings.append(
            f"missing spreadsheet dependency for {path.suffix}: {exc}"
        )

    except Exception as exc:
        warnings.append(
            "spreadsheet parse failed: "
            f"{type(exc).__name__}: {exc}"
        )

    return records, warnings


def parse_file(path: Path, root: Path) -> ParseResult:
    doc = _metadata(path, root)
    if path.suffix.lower() == ".docx":
        return ParseResult(doc, text_evidence=parse_docx(path, doc))
    if path.suffix.lower() == ".doc":
        text, warnings = parse_legacy_doc(path, doc)
        return ParseResult(doc, text_evidence=text, warnings=warnings)
    if path.suffix.lower() == ".pdf":
        try:
            return ParseResult(doc, text_evidence=parse_pdf(path, doc))
        except Exception as exc:
            return ParseResult(doc, warnings=[f"pdf parse failed: {type(exc).__name__}: {exc}"])
    if path.suffix.lower() in {".xlsx", ".xls"}:
        tables, warnings = parse_excel(path, doc)
        return ParseResult(doc, table_evidence=tables, warnings=warnings)
    return ParseResult(doc, warnings=[f"unsupported file type: {path.suffix}"])
